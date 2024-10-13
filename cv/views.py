import os
import uuid
import logging
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from docx import Document
from io import BytesIO
from .models import TempIDModel, PersonalStatement, KeySkill, Education, WorkExperience, WorkTask, Achievement, Interest
from .serializers import (
    PersonalStatementSerializer, KeySkillSerializer, EducationSerializer,
    WorkExperienceSerializer, InterestSerializer
)
import json
from django.views.decorators.csrf import csrf_exempt
import joblib
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer

MODEL_PATH_Statement = os.path.join(os.path.dirname(__file__), 'models', 'statement.joblib')
MODEL_PATH_Skills = os.path.join(os.path.dirname(__file__), 'models', 'skills.joblib')
MODEL_PATH_Interests = os.path.join(os.path.dirname(__file__), 'models', 'interests.joblib')
MODEL_PATH_Education = os.path.join(os.path.dirname(__file__), 'models', 'education.joblib')
MODEL_PATH_Work = os.path.join(os.path.dirname(__file__), 'models', 'work.joblib')
# 加载模型
pipeline_statement = joblib.load(MODEL_PATH_Statement)
pipeline_skills = joblib.load(MODEL_PATH_Skills)
pipeline_interests = joblib.load(MODEL_PATH_Interests)
pipeline_education = joblib.load(MODEL_PATH_Education)
pipeline_work = joblib.load(MODEL_PATH_Work)

MODEL_MAP = {
    'statement': pipeline_statement,  
    'skills': pipeline_skills,
    'interests': pipeline_interests,
    'education': pipeline_education,
    'work': pipeline_work          
}

# 用SVM模型对前端返回的结果进行评估
@csrf_exempt  # 临时禁用CSRF
@api_view(['GET','POST'])
def predict_category(request):
    if request.method == 'POST':
        content = request.POST.get('content', None)
        input_type = request.POST.get('input_type', None)
        print(input_type)
        if content is None or input_type is None:
            return JsonResponse({'error': 'No content provided'}, status=400)
        
        model = MODEL_MAP.get(input_type)
        if model is None:
            return JsonResponse({'error': f'Invalid input_type: {input_type}'}, status=400)
        # 使用模型进行预测
        predicted_category = model.predict([content])
        
        # 返回预测结果
        return JsonResponse({'predicted_category': predicted_category[0]})
    
    return JsonResponse({'error': 'Invalid request method'}, status=400)


# @csrf_exempt
# def submit_info(request):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body)
#             name = data.get('name')

#             if not name:
#                 return JsonResponse({'status': 'error', 'message': 'No user input provided'}, status=400)

#             # 数据处理逻辑，例如保存到数据库
#             return JsonResponse({'status': 'success', 'message': 'Data received successfully'})

#         except json.JSONDecodeError:
#             return JsonResponse({'status': 'error', 'message': 'Invalid JSON format'}, status=400)

#     return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)

# # 获取 logger 对象
# logger = logging.getLogger(__name__)

# @api_view(['GET'])
# def get_all_info_by_temp_id(request, temp_id):
#     try:
#         # Fetch the TempIDModel instance
#         temp_id_instance = TempIDModel.objects.get(temp_id=temp_id)

#         # Fetch related data
#         personal_statement = PersonalStatement.objects.filter(temp_id=temp_id_instance).first()
#         key_skills = KeySkill.objects.filter(temp_id=temp_id_instance)
#         educations = Education.objects.filter(temp_id=temp_id_instance)
#         work_experiences = WorkExperience.objects.filter(temp_id=temp_id_instance)
#         interests = Interest.objects.filter(temp_id=temp_id_instance)

#         # Fetch WorkTasks and Achievements based on related WorkExperience
#         work_experience_ids = work_experiences.values_list('id', flat=True)
#         work_tasks = WorkTask.objects.filter(work_experience__id__in=work_experience_ids)
#         achievements = Achievement.objects.filter(work_experience__id__in=work_experience_ids)

#         # Load the Word template document
#         template_path = os.path.join(settings.BASE_DIR, 'documents', str(temp_id) + '.docx')  # Ensure this file exists
#         doc = Document()
#         doc.save(template_path)
#         if not os.path.exists(template_path):
#             return Response({'error': f'Template file not found at {template_path}'}, status=status.HTTP_404_NOT_FOUND)

#         doc = Document(template_path)

#         # Modify the document by adding the required information
#         doc.add_heading('Name', level=1)  # Adding Name as a heading
#         doc.add_paragraph('[Student Name]')  # Replace with dynamic student name


#         # Add Personal Statement
#         if personal_statement:
#             doc.add_heading('Personal Statement', level=1)
#             doc.add_paragraph(personal_statement.content)

#         # Add Key Skills
#         if key_skills.exists():
#             doc.add_heading('Key Skills', level=1)
#             for skill in key_skills:
#                 doc.add_paragraph(skill.content, style='ListBullet')

#         # Add Education
#         if educations.exists():
#             doc.add_heading('Education', level=1)
#             for education in educations:
#                 doc.add_paragraph(f"{education.programme_title} ({education.year_start} - {education.year_complete})")
#                 doc.add_paragraph(f"Institution: {education.institution}, {education.location}")
#                 if education.courses_and_projects:
#                     doc.add_paragraph(f"Courses/Projects: {education.courses_and_projects}")

#         # Add Work Experience
#         if work_experiences.exists():
#             doc.add_heading('Work Experience', level=1)
#             for experience in work_experiences:
#                 doc.add_paragraph(f"{experience.job_title} ({experience.year_start} - {experience.year_complete})")
#                 doc.add_paragraph(f"Organisation: {experience.organisation}, {experience.location}")

#                 # Add Work Tasks
#                 related_tasks = work_tasks.filter(work_experience=experience)
#                 if related_tasks.exists():
#                     doc.add_heading('Tasks', level=2)
#                     for task in related_tasks:
#                         doc.add_paragraph(task.content, style='ListBullet')

#                 # Add Achievements
#                 related_achievements = achievements.filter(work_experience=experience)
#                 if related_achievements.exists():
#                     doc.add_heading('Achievements', level=2)
#                     for achievement in related_achievements:
#                         doc.add_paragraph(achievement.content, style='ListBullet')

#         # Add Interests
#         if interests.exists():
#             doc.add_heading('Interests', level=1)
#             for interest in interests:
#                 doc.add_paragraph(interest.content)

#         doc.add_heading('Reference', level=1)  # Adding Reference as a heading
#         doc.add_paragraph('[Reference Details]')  # Replace with dynamic reference

#         # Define the save path for the new document
#         save_path = os.path.join(settings.BASE_DIR, 'documents', f"{temp_id}.docx")
#         print(f"Attempting to save document at: {save_path}")  # Debugging output

#         # Save the document to disk
#         try:
#             doc.save(save_path)
#             print(f"Document successfully saved at: {save_path}")
#         except Exception as e:
#             print(f"Error saving document: {e}")
#             return Response({'error': f"Error saving document: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#         # Check if the document was saved successfully
#         if not os.path.exists(save_path):
#             return Response({'error': 'Document was not saved as expected.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#         # Save document to BytesIO for response
#         byte_io = BytesIO()
#         doc.save(byte_io)
#         byte_io.seek(0)

#         # Create a response with the document
#         response = HttpResponse(byte_io,
#                                 content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = f'attachment; filename="{temp_id}.docx"'
#         return response

#     except TempIDModel.DoesNotExist:
#         return Response({'error': 'TempID not found'}, status=status.HTTP_404_NOT_FOUND)
    
# @api_view(['GET'])
# def generate_temp_id(request):
#     temp_id_instance = TempIDModel.objects.create()  # 创建新的 TempIDModel 实例，这将生成 temp_id
#     return Response({'temp_id': str(temp_id_instance.temp_id)}, status=status.HTTP_200_OK)

# @api_view(['POST'])
# def submit_all_info(request):
    # print(request.body)  # 打印原始请求体
    # data = json.loads(request.body)
    # print(data)  # 打印解析后的数据，检查其是否为字典
    
    try:
        data = request.data
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except json.JSONDecodeError as e:
                return Response({'error': 'Invalid JSON format'}, status=status.HTTP_400_BAD_REQUEST)

        # 确认数据是否为字典
        if not isinstance(data, dict):
            return Response({'error': 'Expected data to be a dictionary'}, status=status.HTTP_400_BAD_REQUEST)


        logger.info(f"Received data: {data}")

        # 获取前端传递的 Temp ID
        temp_id = data.get('temp_id')
        if not temp_id:
            logger.error("Temp ID is missing in the request")
            return Response({'error': 'Temp ID is required'}, status=status.HTTP_400_BAD_REQUEST)


        temp_id_instance = TempIDModel.objects.get(temp_id=temp_id)
        
        # 保存 Personal Statement
        personal_statement_data = data.get('personal_statement')
        if personal_statement_data:
            personal_statement_instance = PersonalStatement.objects.create(
                temp_id=temp_id_instance,
                content=f"Q1: {personal_statement_data.get('statementQ1', '')}\nQ2: {personal_statement_data.get('statementQ2', '')}"
            )
        
        # 保存 Key Skills
        key_skills_data = data.get('key_skills', [])
        for skill_data in key_skills_data:
            KeySkill.objects.create(
                temp_id=temp_id_instance,
                content=skill_data.get('content', '')
            )

        # 保存 Education
        education_data = data.get('educations', [])
        for edu_data in education_data:
            Education.objects.create(
                temp_id=temp_id_instance,
                programme_title=edu_data.get('programme_title', ''),
                institution=edu_data.get('institution', ''),
                location=edu_data.get('location', ''),
                year_start=edu_data.get('year_start', ''),
                year_complete=edu_data.get('year_complete', ''),
                courses_and_projects=edu_data.get('courses_and_projects', '')
            )

        # 保存 Work Experience
        work_experiences_data = data.get('work_experiences', [])
        for work_data in work_experiences_data:
            work_experience_instance = WorkExperience.objects.create(
                temp_id=temp_id_instance,
                job_title=work_data.get('job_title', ''),
                organisation=work_data.get('organisation', ''),
                location=work_data.get('location', ''),
                year_start=work_data.get('year_start', ''),
                year_complete=work_data.get('year_complete', '')
            )
            
            # 保存 Work Tasks
            work_tasks_data = work_data.get('tasks', [])
            for task_content in work_tasks_data:
                WorkTask.objects.create(
                    work_experience=work_experience_instance,
                    content=task_content
                )
            
            # 保存 Achievements
            achievements_data = work_data.get('achievements', [])
            for achievement_content in achievements_data:
                Achievement.objects.create(
                    work_experience=work_experience_instance,
                    content=achievement_content
                )

        # 保存 Interests
        interests_data = data.get('interests', [])
        for interest_data in interests_data:
            Interest.objects.create(
                temp_id=temp_id_instance,
                content=interest_data.get('content', '')
            )


        # 加载模板文档
        template_path = os.path.join(settings.BASE_DIR, 'documents', 'template.docx')  # 确保这个文件存在
        if not os.path.exists(template_path):
            logger.error(f"Template file not found at {template_path}")
            return Response({'error': f'Template file not found at {template_path}'}, status=status.HTTP_404_NOT_FOUND)

        # 打开模板文档
        document = Document(template_path)

        # 修改文档内容
        document.add_heading('User Submitted Data', level=1)

        # 添加个人陈述部分
        if 'personal_statement' in data:
            document.add_heading('Personal Statement', level=2)
            document.add_paragraph(f"Statement Q1: {data['personal_statement'].get('statementQ1', '')}")
            document.add_paragraph(f"Statement Q2: {data['personal_statement'].get('statementQ2', '')}")

        # 添加关键技能部分
        if 'key_skills' in data:
            document.add_heading('Key Skills', level=2)
            for skill in data['key_skills']:
                document.add_paragraph(skill.get('content', ''), style='ListBullet')

        # 添加教育背景部分
        if 'educations' in data:
            document.add_heading('Education', level=2)
            for education in data['educations']:
                document.add_paragraph(f"Major: {education.get('major', '')}")
                document.add_paragraph(f"School: {education.get('school', '')}")
                document.add_paragraph(f"Start Time: {education.get('startTime', '')}")
                document.add_paragraph(f"End Time: {education.get('endTime', '')}")
                document.add_paragraph(f"Achievements: {education.get('achievements', '')}")

        # 添加工作经验部分
        if 'work_experiences' in data:
            document.add_heading('Work Experience', level=2)
            for work in data['work_experiences']:
                document.add_paragraph(f"Job Title: {work.get('job_title', '')}")
                document.add_paragraph(f"Organisation: {work.get('organisation', '')}")
                document.add_paragraph(f"Start Time: {work.get('startTime', '')}")
                document.add_paragraph(f"End Time: {work.get('endTime', '')}")
                document.add_paragraph(f"Tasks: {work.get('tasks', '')}", style='ListBullet')
                document.add_paragraph(f"Achievements: {work.get('achievements', '')}", style='ListBullet')

        # 添加兴趣部分
        if 'interests' in data:
            document.add_heading('Interests', level=2)
            for interest in data['interests']:
                document.add_paragraph(interest.get('content', ''))

        # 保存文档到指定路径
        output_path = os.path.join(settings.BASE_DIR, 'documents', f'submitted_data_{temp_id}.docx')
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        document.save(output_path)
        logger.info(f"Document saved successfully at {output_path}")

        # 将文档保存到 BytesIO 对象中，以便于返回响应
        byte_io = BytesIO()
        document.save(byte_io)
        byte_io.seek(0)

        # 创建响应并附加文档
        response = HttpResponse(byte_io,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{temp_id}.docx"'
        return Response({'status': 'success'}, status=status.HTTP_200_OK)
    
    except TempIDModel.DoesNotExist:
        return Response({'error': 'TempID not found'}, status=status.HTTP_404_NOT_FOUND)
 
    except Exception as e:
        logger.error(f"An error occurred: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def add_horizontal_line(paragraph):
    """在段落中添加一条横线"""
    p = paragraph._p  # 获取底层段落对象
    pPr = p.get_or_add_pPr()  # 获取段落属性

    # 创建 <w:pBdr> 元素并设置下边框
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # 线类型: 'single' 是单线
    bottom.set(qn('w:sz'), '6')  # 线宽度
    bottom.set(qn('w:space'), '1')  # 与内容的间距
    bottom.set(qn('w:color'), 'auto')  # 颜色: 'auto' 是自动黑色

    pBdr.append(bottom)
    pPr.append(pBdr)
    
    
@csrf_exempt  # 临时禁用CSRF
@api_view(['GET','POST'])
def generate_word(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))  # 先将 request.body 解码为字符串，然后加载为 JSON
        except json.JSONDecodeError:
            return HttpResponse("Invalid JSON", status=400)
        
        # return HttpResponse(request.body.decode('utf-8'))
        document = Document()
        
        # PersonalInof
        personalinfo_raw = data["PersonalInfo"]
        personalinfo = json.loads(personalinfo_raw)
        # document.add_heading('Personal Information', level=1)
        document.add_heading('Name  ' + personalinfo["name"], level=1) 
        add_horizontal_line(document.add_paragraph()) #horizontal_line
        document.add_paragraph('Mobile:   ' + personalinfo["phone"])  
        document.add_paragraph('Email:    ' + personalinfo["email"]) 
        document.add_paragraph('LinkedIn URL:   ' + personalinfo["linkedin"]) 
                
        # # PersonalStatement
        personalstatement_raw = data["PersonalStatement"]
        personalstatement = json.loads(personalstatement_raw)
        document.add_heading('Personal Statement OR Career Objective', level=1)        
        add_horizontal_line(document.add_paragraph()) #horizontal_line
        document.add_paragraph(personalstatement["Q1"])
        document.add_paragraph(personalstatement["Q2"])
        
        # Key Skills
        skills_raw = data["Skills"]
        skills = json.loads(skills_raw)
        document.add_heading('Skills', level=1)        
        add_horizontal_line(document.add_paragraph()) #horizontal_line
        document.add_paragraph(skills["Q1"])
        document.add_paragraph(skills["Q2"])
        
        # Education
        education_raw = data["Education"]
        education_data = json.loads(education_raw)
        document.add_heading('Education', level=1)      
        for education in education_data.values():
            add_horizontal_line(document.add_paragraph()) #horizontal_line
            table = document.add_table(rows=1, cols=2)
            table.columns[0].width = Pt(400)
            table.columns[1].width = Pt(400)
            cell_left = table.cell(0, 0)
            cell_left.text = education.get('major','N/A')
            cell_right = table.cell(0, 1)
            paragraph = cell_right.paragraphs[0]
            run = paragraph.add_run(f"dates({education.get('start','N/A')}-{education.get('end','N/A')})")
            paragraph.alignment = 2  # 2 表示右对齐
            document.add_paragraph(education["school"])            
            document.add_paragraph(education["achievements"])
            
            
        # Work Experience
        work_raw = data["Work"]
        work_data = json.loads(work_raw)
        document.add_heading('Work', level=1)      
        for work in work_data.values():
            add_horizontal_line(document.add_paragraph()) #horizontal_line
            table = document.add_table(rows=1, cols=2)
            table.columns[0].width = Pt(400)
            table.columns[1].width = Pt(400)
            cell_left = table.cell(0, 0)
            cell_left.text = work.get('job_title','N/A')
            cell_right = table.cell(0, 1)
            paragraph = cell_right.paragraphs[0]
            run = paragraph.add_run(f"dates({work.get('startTime','N/A')}-{work.get('endTime','N/A')})")
            paragraph.alignment = 2  # 2 表示右对齐
            document.add_paragraph(work["organisation"])            
            document.add_paragraph(work["tasks"])      
            document.add_paragraph(work["achievements"])
        
        # # Interests
        interests_raw = data["Interests"]
        interests = json.loads(interests_raw)
        document.add_heading('Interests', level=1)        
        add_horizontal_line(document.add_paragraph()) #horizontal_line
        document.add_paragraph(interests["Q1"])
        document.add_paragraph(interests["Q2"])
        
        
        
        
        

        # # 将 Word 文档保存到内存中的字节流
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=PersonalStatement.docx'

        # # 将文档写入到响应中
        document.save(response)

        return response
    else:
        return HttpResponse(status=400)